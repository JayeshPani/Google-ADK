"""Resume file parsing and normalization."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
import unicodedata


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


@dataclass(slots=True)
class ParsedResume:
    file_path: Path
    file_name: str
    file_type: str
    raw_text: str
    normalized_text: str
    ats_findings: list[str] = field(default_factory=list)
    contact_signals: dict[str, bool] = field(default_factory=dict)


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text).replace("\xa0", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_pdf(path: Path) -> str:
    import fitz

    with fitz.open(path) as document:
        return "\n".join(page.get_text("text") for page in document)


def _read_docx(path: Path) -> tuple[str, list[str]]:
    from docx import Document

    document = Document(path)
    table_warnings: list[str] = []
    if document.tables:
        table_warnings.append("Resume contains tables; ATS parsers often flatten them poorly.")
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraph_text), table_warnings


def _read_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _detect_contact_signals(text: str) -> dict[str, bool]:
    lowered = text.lower()
    return {
        "email": bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)),
        "phone": bool(re.search(r"(\+\d{1,3}\s?)?[\(\- ]?\d{3}[\)\- ]?\d{3}[\- ]?\d{4}", text)),
        "linkedin": "linkedin.com" in lowered,
        "github": "github.com" in lowered,
        "portfolio": any(token in lowered for token in ("portfolio", "website", "https://")),
    }


def _detect_ats_findings(text: str, contact_signals: dict[str, bool]) -> list[str]:
    findings: list[str] = []
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not contact_signals["email"]:
        findings.append("No email detected in the resume header.")
    if not contact_signals["linkedin"]:
        findings.append("No LinkedIn URL detected; many recruiters look for it on student resumes.")
    if max((len(line.split()) for line in lines), default=0) > 40:
        findings.append("At least one bullet is very long; ATS and recruiters both prefer tighter bullets.")
    if text.count("•") + text.count("- ") < 4:
        findings.append("The resume has few explicit bullets; key evidence may be hard to scan quickly.")
    if re.search(r"[●◆■★]", text):
        findings.append("Decorative bullet symbols detected; plain bullets are safer for ATS parsing.")
    if len(lines) < 8:
        findings.append("Very little parsed text was extracted; formatting may not be ATS-friendly.")
    return findings


def parse_resume_file(file_path: str | Path) -> ParsedResume:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported resume format: {suffix or 'unknown'}")

    ats_findings: list[str] = []
    if suffix == ".pdf":
        raw_text = _read_pdf(path)
    elif suffix == ".docx":
        raw_text, table_warnings = _read_docx(path)
        ats_findings.extend(table_warnings)
    else:
        raw_text = _read_plain_text(path)

    normalized = normalize_text(raw_text)
    contact_signals = _detect_contact_signals(normalized)
    ats_findings.extend(_detect_ats_findings(normalized, contact_signals))
    return ParsedResume(
        file_path=path,
        file_name=path.name,
        file_type=suffix.lstrip("."),
        raw_text=raw_text,
        normalized_text=normalized,
        ats_findings=ats_findings,
        contact_signals=contact_signals,
    )

