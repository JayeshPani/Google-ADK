"""Binary export helpers for rewritten resumes."""

from __future__ import annotations

from io import BytesIO

from job_rejection_agent.domain import RewrittenResume, ResumeSectionBlock


def _iter_sections(resume: RewrittenResume) -> list[ResumeSectionBlock]:
    return [
        resume.summary,
        resume.experience,
        resume.projects,
        resume.skills,
        resume.education,
    ]


def build_resume_docx_bytes(resume: RewrittenResume) -> bytes:
    from docx import Document

    document = Document()
    header_lines = resume.header.items
    if header_lines:
        document.add_heading(header_lines[0], level=0)
        for line in header_lines[1:]:
            document.add_paragraph(line)
    for section in _iter_sections(resume):
        document.add_heading(section.title, level=1)
        if section.note:
            document.add_paragraph(section.note)
        for item in section.items:
            document.add_paragraph(item, style="List Bullet")
    if resume.ats_notes:
        document.add_heading("ATS Notes", level=1)
        for note in resume.ats_notes:
            document.add_paragraph(note, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_resume_pdf_bytes(resume: RewrittenResume) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=LETTER, leftMargin=48, rightMargin=48, topMargin=48, bottomMargin=48)
    styles = getSampleStyleSheet()
    story = []

    header_lines = resume.header.items
    if header_lines:
        story.append(Paragraph(header_lines[0], styles["Title"]))
        for line in header_lines[1:]:
            story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 12))

    for section in _iter_sections(resume):
        story.append(Paragraph(section.title, styles["Heading2"]))
        if section.note:
            story.append(Paragraph(section.note, styles["Italic"]))
        items = [ListItem(Paragraph(item, styles["BodyText"])) for item in section.items]
        if items:
            story.append(ListFlowable(items, bulletType="bullet"))
        else:
            story.append(Paragraph("No verified content available for this section yet.", styles["BodyText"]))
        story.append(Spacer(1, 10))

    if resume.ats_notes:
        story.append(Paragraph("ATS Notes", styles["Heading2"]))
        story.append(ListFlowable([ListItem(Paragraph(item, styles["BodyText"])) for item in resume.ats_notes], bulletType="bullet"))

    document.build(story)
    return buffer.getvalue()
